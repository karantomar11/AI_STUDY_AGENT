import os
from openai import OpenAI
from docx import Document
from docx.shared import Pt as Points
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_feynman_doc(master_context, api_key, output_path):
    """
    Generates a Feynman Mastery Page using OpenRouter (MiMo-v2-Flash) and saves it as a .docx file.
    """
    print("\n--- [Phase 5] Feynman Mastery Module Initialized ---")
    print("Connecting to OpenRouter (Model: xiaomi/mimo-v2-flash:free) with Reasoning Enabled...")

    client = OpenAI(
        base_url=os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1"),
        api_key=api_key
    )

    prompt = """
    ACT AS: A World-Class Science Communicator (like Richard Feynman or Steve Blue).
TASK: Synthesize the 'Master Context' into a final 'Feynman Mastery Page.'

GOAL: Explain the topic so perfectly that a 12-year-old can explain it to their friends, but a Senior Engineer can still find the logic sound.

STRICT OUTPUT SCHEMA:
# FEYNMAN MASTERY: [CONCEPT NAME]

## I. The One-Sentence Essence
- Provide a single, jargon-free sentence that captures the "Soul" of this topic. No words longer than 3 syllables.

## II. The "Playground" Analogy (The 12-Year-Old Level)
- **The Setup:** Describe a common scenario (Video games, LEGO, baking, or school sports).
- **The Execution:** Map the technical logic of the notes to the rules of the playground.
- **The Connection:** Explicitly state: "Just like [Analogy Part], Software Engineering uses [Technical Concept] to..."

## III. The Jargon Decrypter
- Provide a 3-column table: | Term | Academic Definition | "Human" Translation |
- Include at least 5 critical terms from the notes.

## IV. The "Blind Spot" Audit (The Deep Reasoner)
- Identify 3 areas where people usually get this wrong. 
- Explain why they get it wrong and what the "Aha!" moment is.

## V. The Real-World "So What?"
- Describe one scenario where knowing this saved a project or one where NOT knowing it caused a disaster.
    context:
    ---
    {context}
    ---
    """.format(context=master_context)

    try:
        completion = client.chat.completions.create(
            model="xiaomi/mimo-v2-flash:free",
            messages=[
                {"role": "user", "content": prompt}
            ],
            extra_body={"reasoning_enabled": True} # Activate Reasoning Engine
        )
        
        response_text = completion.choices[0].message.content
        
        if not response_text:
            print("Error: Empty response from AI.")
            return False

        # --- Document Creation ---
        print(f"Formatting Feynman Mastery Page to {output_path}...")
        doc = Document()
        
        # Title
        title = doc.add_heading('Feynman Mastery Method', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add content, attempting to bold headers if they appear in the text
        # Since the AI output can vary, we will dump the text but try to style it nicely.
        # A stronger approach is to parse the 4 sections, but for now we will format the raw text 
        # and ensure the file exists as requested.
        
        # Let's split by lines to handle some basic formatting
        lines = response_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Simple heuristic for headers based on the prompt structure
            if any(header in line for header in ["The Core Essence", "The 12-Year-Old Analogy", "The Jargon Translator", "The 'So What?'"]):
                p = doc.add_heading(line, level=2)
            else:
                p = doc.add_paragraph(line)
        
        doc.save(output_path)
        print(f"✅ Saved Phase 5 Output: {output_path}")
        
        # --- Verification Print ---
        # "Print the '12-Year-Old Analogy' to the terminal to confirm the logic quality."
        print("\n--- VERIFICATION: 12-Year-Old Analogy ---")
        found_analogy = False
        capture = False
        analogy_text = []
        
        for line in lines:
            if "The 12-Year-Old Analogy" in line:
                capture = True
                found_analogy = True
                continue # Skip the header itself
            if capture:
                # Stop if we hit the next header
                if any(header in line for header in ["The Jargon Translator", "The 'So What?'"]):
                    capture = False
                    break
                analogy_text.append(line)
        
        if found_analogy:
            print("\n".join(analogy_text).strip())
        else:
            print("(Analogy section not strictly found in output, printing first 500 chars of response instead)")
            print(response_text[:500] + "...")
            
        return True

    except Exception as e:
        print(f"Error in Feynman Module: {e}")
        return False

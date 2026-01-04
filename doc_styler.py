import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
try:
    from math2docx import add_math
except ImportError:
    add_math = None
    print("Warning: math2docx not installed. Math equations will not be rendered natively.")

class DocStyler:
    def __init__(self):
        # Professional Dark Blue: RGB(0, 51, 102)
        self.header_color = RGBColor(0, 51, 102)

    def apply_rich_styling(self, paragraph, text):
        """
        PERMANENT FIX: Converts $...$ into real Word Equations.
        Parses a string for Bold, Italic, Underline, Highlight, and Math.
        """
        # Pattern to find math, bold, highlights, underline, italic
        # Order: Math ($), Bold (**), Highlight (==), Underline (__), Italic (*)
        pattern = r'(\$.*?\$|\*\*.*?\*\*|==.*?==|__.*?__|(?<!\*)\*(?!\*))'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            # --- MATH ($) ---
            if part.startswith('$') and part.endswith('$'):
                content = part[1:-1]
                # ADVANCED FIX: Native Word Math
                # We assume math2docx is installed as per requirements
                if add_math:
                    try:
                        add_math(paragraph, content)
                    except Exception as e:
                        # PERMANENT RELIABILITY FIX: Fallback if LaTeX is invalid
                        print(f"⚠️ Math Render Warning: Could not render '{content}' due to: {e}. Falling back to text.")
                        paragraph.add_run(part)
                else:
                    paragraph.add_run(part) # Fallback if import failed
            
            # --- BOLD (**) ---
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                
            # --- HIGHLIGHT (==) ---
            elif part.startswith('==') and part.endswith('=='):
                run = paragraph.add_run(part[2:-2])
                # User requested Color Index 4 (Turquoise/Bright Green)
                run.font.highlight_color = 4
                
            # --- UNDERLINE (__) ---
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.underline = True
            
            # --- ITALIC (*) ---
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                
            # --- NORMAL TOPIC ---
            else:
                paragraph.add_run(part)

    def create_styled_doc(self, markdown_text, output_path, title="Study Document"):
        doc = Document()
        
        # Add Title (Level 0)
        t = doc.add_heading(title, level=0)
        t.alignment = 1 # Center it
        
        lines = markdown_text.split('\n')
        
        in_table = False
        table_data = []
        
        for line in lines:
            stripped = line.strip()
            
            # --- TABLE DETECTION ---
            if stripped.startswith('|') and stripped.endswith('|'):
                 # Check if it's a separator row
                 if set(stripped.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')) == set():
                     continue 
                 in_table = True
                 table_data.append([c.strip() for c in stripped.split('|')[1:-1]])
                 continue
            else:
                if in_table:
                    # Render Table
                    if table_data:
                        rows = len(table_data)
                        cols = len(table_data[0]) if rows > 0 else 0
                        if rows > 0:
                            tbl = doc.add_table(rows=rows, cols=cols)
                            tbl.style = 'Table Grid'
                            for r_idx, r_dat in enumerate(table_data):
                                row = tbl.rows[r_idx]
                                for c_idx, cell_text in enumerate(r_dat):
                                    if c_idx < len(row.cells):
                                        cell = row.cells[c_idx]
                                        cell._element.clear_content()
                                        p = cell.add_paragraph()
                                        self.apply_rich_styling(p, cell_text)
                    in_table = False
                    table_data = []

            if not stripped: continue
            
            # --- HEADERS ---
            if stripped.startswith('# '):
                h = doc.add_heading(stripped[2:], level=1)
                for r in h.runs: r.font.color.rgb = self.header_color
            elif stripped.startswith('## '):
                h = doc.add_heading(stripped[3:], level=2)
                for r in h.runs: r.font.color.rgb = self.header_color
            elif stripped.startswith('### '):
                h = doc.add_heading(stripped[4:], level=3)
                # Level 3 default is fine
            
            # --- LISTS ---
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                self.apply_rich_styling(p, stripped[2:])
            
            # --- NORMAL ---
            else:
                p = doc.add_paragraph()
                self.apply_rich_styling(p, stripped)

        doc.save(output_path)
        print(f"✅ Styled Document Saved: {output_path}")

# Helper for existing calls
def create_styled_docx(markdown_text, output_path, title="Study Notes"):
    styler = DocStyler()
    # If title kwarg is not effective in create_styled_doc usage loop (it is handled), this is fine.
    # We should extract a title from markdown if possible, but default is okay.
    # The previous prompts put "# TITLE" in the markdown, so create_styled_doc will see it as H1 usually.
    # But strictly, the first line logic in create_styled_doc assumes '# ' is H1.
    # If the markdown has a title, it will be rendered.
    styler.create_styled_doc(markdown_text, output_path, title)
